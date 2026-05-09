import docx
import json
import os

def check_docx():
    file_path = r"d:\Downloads\Therapy Merged\Guided Series.docx"
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return
    
    doc = docx.Document(file_path)
    print(f"Number of tables: {len(doc.tables)}")
    
    for i, table in enumerate(doc.tables):
        print(f"\nTable {i}:")
        for r_idx, row in enumerate(table.rows):
            if r_idx < 10: # Only first 10 rows
                print([cell.text.strip() for cell in row.cells])
            else:
                print("...")
                break

check_docx()
