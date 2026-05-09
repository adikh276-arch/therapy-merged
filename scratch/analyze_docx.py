import docx
import os
import json

def read_guided_series():
    file_path = r"d:\Downloads\Therapy Merged\Guided Series.docx"
    doc = docx.Document(file_path)
    
    # Assuming the first table contains the "all concerns" data
    if not doc.tables:
        print("No tables found in docx")
        return

    table = doc.tables[0]
    data = []
    for row in table.rows:
        data.append([cell.text.strip() for cell in row.cells])
    
    return data

def analyze_data(table_data):
    if not table_data:
        return
    
    headers = table_data[0]
    rows = table_data[1:]
    
    # Group by Concern (assuming first column is Concern)
    concerns = {}
    for row in rows:
        if not row: continue
        concern = row[0]
        if concern not in concerns:
            concerns[concern] = []
        
        # Collect activities from other columns
        # Some rows might be activities for the same concern
        # Or each row is a new activity?
        # Let's just collect all non-empty strings in the row except the first column
        activities = [item for item in row[1:] if item and item.strip()]
        concerns[concern].extend(activities)
    
    summary = {}
    for concern, activities in concerns.items():
        # Count unique activities
        unique_activities = list(set(activities))
        # Filter out generic terms like "Activity", "Description" if they appear as values
        unique_activities = [a for a in unique_activities if a.lower() not in ["activity", "description", "details", "concern"]]
        
        # Check for "missing" or empty-looking placeholders
        missing_count = sum(1 for a in unique_activities if "missing" in a.lower() or "tbd" in a.lower() or "to be added" in a.lower())
        
        summary[concern] = {
            "total_activities": len(unique_activities),
            "missing_count": missing_count,
            "activities": unique_activities
        }
    
    return summary

data = read_guided_series()
if data:
    # Print the raw first 5 rows to see structure
    print("RAW ROWS (first 5):")
    for r in data[:5]:
        print(r)
    
    summary = analyze_data(data)
    print("\nSUMMARY:")
    print(json.dumps(summary, indent=2))
