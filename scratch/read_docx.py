import docx
import json
import os

def extract_docx_content(file_path):
    doc = docx.Document(file_path)
    content = []
    
    # Extract tables
    for i, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        content.append({"type": "table", "index": i, "data": table_data})
    
    # Extract paragraphs (if needed)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    content.append({"type": "paragraphs", "data": paragraphs})
    
    return content

file_path = r"d:\Downloads\Therapy Merged\Guided Series.docx"
if os.path.exists(file_path):
    try:
        data = extract_docx_content(file_path)
        print(json.dumps(data, indent=2))
    except Exception as e:
        print(f"Error: {e}")
else:
    print(f"File not found: {file_path}")
